from docx import Document
from datetime import date
from io import BytesIO
import re

MESES_ES = {
    1: 'enero', 2: 'febrero', 3: 'marzo',
    4: 'abril', 5: 'mayo', 6: 'junio',
    7: 'julio', 8: 'agosto', 9: 'septiembre',
    10: 'octubre', 11: 'noviembre', 12: 'diciembre'
}


def construir_contexto(usuario, categoria_plantilla=None):
    """
    Construye el diccionario de marcadores según el rol del usuario autenticado
    y su proyecto. Los datos se completan según la categoría de la plantilla.
    """
    from apps.proyectos.models import ProyectoEstudiante

    hoy = date.today()
    roles = list(usuario.roles.values_list('rol', flat=True))

    contexto = {
        'dia': str(hoy.day),
        'mes': MESES_ES[hoy.month],
        'anio': str(hoy.year),
        'fecha_actual': f"{hoy.day} de {MESES_ES[hoy.month]} de {hoy.year}",
        'nombre_departamento': 'Departamento de Sistemas',
        'nombre_universidad': 'Universidad de Nariño',
        'cedula_asesor': '',
        'nombre_asesor': '',
        'correo_asesor': '',
        'celular_asesor': '',
        'cedula_coasesor': '',
        'nombre_coasesor': '',
        'correo_coasesor': '',
        'celular_coasesor': '',
        'codigo_estudiante': '',
        'nombre_estudiante': '',
        'correo_estudiante': '',
        'celular_estudiante': '',
        'cargo_docente': 'Docente',
        'cargo_jurado': 'Jurado Evaluador',
    }

    # Para plantillas de estudiante - usar datos del proyecto
    if categoria_plantilla == 'Estudiante' or 'ESTUDIANTE' in roles:
        try:
            pe = ProyectoEstudiante.objects.select_related(
                'proyecto__asesor__usuario',
                'proyecto__coasesor__usuario',
                'estudiante__usuario'
            ).get(estudiante__usuario=usuario)

            proyecto = pe.proyecto
            estudiante = pe.estudiante
            asesor = proyecto.asesor
            coasesor = getattr(proyecto, 'coasesor', None)

            contexto.update({
                'codigo_estudiante': estudiante.codigo_estudiante,
                'nombre_estudiante': f"{usuario.nombre} {usuario.apellido}",
                'correo_estudiante': usuario.correo,
                'celular_estudiante': getattr(usuario, 'celular', ''),
            })

            if asesor and hasattr(asesor, 'usuario'):
                au = asesor.usuario
                contexto.update({
                    'cedula_asesor': getattr(au, 'numero_documento', ''),
                    'nombre_asesor': f"{au.nombre} {au.apellido}",
                    'correo_asesor': au.correo,
                    'celular_asesor': getattr(au, 'celular', ''),
                })

            if coasesor and hasattr(coasesor, 'usuario'):
                cu = coasesor.usuario
                contexto.update({
                    'cedula_coasesor': getattr(cu, 'numero_documento', ''),
                    'nombre_coasesor': f"{cu.nombre} {cu.apellido}",
                    'correo_coasesor': cu.correo,
                    'celular_coasesor': getattr(cu, 'celular', ''),
                })

        except ProyectoEstudiante.DoesNotExist:
            pass

    # Para docente/asesor - usar sus propios datos
    if 'DOCENTE' in roles:
        contexto.update({
            'cedula_asesor': getattr(usuario, 'numero_documento', ''),
            'nombre_asesor': f"{usuario.nombre} {usuario.apellido}",
            'correo_asesor': usuario.correo,
            'celular_asesor': getattr(usuario, 'celular', ''),
        })

    # Para jurado - usar sus propios datos
    if 'JURADO' in roles:
        contexto.update({
            'cedula_asesor': getattr(usuario, 'numero_documento', ''),
            'nombre_asesor': f"{usuario.nombre} {usuario.apellido}",
            'correo_asesor': usuario.correo,
            'celular_asesor': getattr(usuario, 'celular', ''),
        })

    return contexto


def reemplazar_en_elemento(elemento, contexto):
    """
    Reemplaza marcadores {{clave}} en un elemento (párrafo o celda).
    Maneja correctamente múltiples runs preservando el formato.
    """
    texto = elemento.text
    if not texto:
        return

    # Buscar si hay marcadores en el texto
    patron = re.compile(r'\{\{(\w+)\}\}')
    matches = patron.findall(texto)

    if not matches:
        return

    # Reemplazar los marcadores encontrados
    for clave in matches:
        valor = contexto.get(clave, '')
        patron_clave = re.compile(r'\{\{' + clave + r'\}\}')
        texto = patron_clave.sub(str(valor), texto)

    # Limpiar runs existentes y crear uno nuevo con el texto modificado
    if hasattr(elemento, 'clear'):
        # es un párrafo
        elemento.clear()
        nuevo_run = elemento.add_run(texto)
    else:
        # es una celda
        for run in list(elemento.runs):
            elemento._element.remove(run._element)
        nuevo_run = elemento.add_run(texto)


def procesar_documento(doc, contexto):
    """
    Procesa todos los elementos del documento:
    - párrafos principales
    - tablas y sus celdas
    - headers y footers
    """
    # Párrafos principales
    for parrafo in doc.paragraphs:
        reemplazar_en_elemento(parrafo, contexto)

    # Tablas
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for parrafo in celda.paragraphs:
                    reemplazar_en_elemento(parrafo, contexto)

    # Headers
    for seccion in doc.sections:
        for parrafo in seccion.header.paragraphs:
            reemplazar_en_elemento(parrafo, contexto)

        for parrafo in seccion.header.paragraphs:
            reemplazar_en_elemento(parrafo, contexto)

    # Footers
    for seccion in doc.sections:
        for parrafo in seccion.footer.paragraphs:
            reemplazar_en_elemento(parrafo, contexto)


def generar_docx_prellenado(plantilla_obj, usuario):
    """
    Toma el archivo .docx de la plantilla,
    reemplaza todos los marcadores con datos reales
    y retorna un BytesIO listo para descargar.
    """
    # Obtener categoría de la plantilla
    categoria = getattr(plantilla_obj, 'categoria', None)
    contexto = construir_contexto(usuario, categoria)

    # Abrir documento
    doc = Document(plantilla_obj.archivo_url.path)

    # Procesar documento
    procesar_documento(doc, contexto)

    # Guardar en buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer